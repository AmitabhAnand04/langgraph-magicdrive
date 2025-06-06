import os
from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, StorageContext, load_index_from_storage
from llama_index.core import Settings
from llama_index.llms.gemini import Gemini
from llama_index.embeddings.gemini import GeminiEmbedding
from llama_index.core.schema import Document
from llama_index.core.memory import ChatMemoryBuffer
from azure.storage.blob import ContainerClient
import shutil
from io import BytesIO
from docx import Document as DocxDocument


api_key = os.getenv("GOOGLE_API_KEY")
# index_storage_dir = "index_storage"
# data_dir = "data"
current_dir = os.path.dirname(os.path.abspath(__file__))  # this gives you .../tools/kb_tools
index_storage_dir = os.path.join(current_dir, "feature_indexing")
data_dir = os.path.join(current_dir, "data")
safety_settings = [
    {
        "category": "HARM_CATEGORY_DANGEROUS",
        "threshold": "BLOCK_LOW_AND_ABOVE",
    },
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_LOW_AND_ABOVE",
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_LOW_AND_ABOVE",
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_LOW_AND_ABOVE",
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_LOW_AND_ABOVE",
    },
]

Settings.embed_model = GeminiEmbedding(
    model_name=os.getenv('RAG_EMBEDDING_MODEL'), api_key=api_key
)

Settings.llm = Gemini(model_name=os.getenv('RAG_LLM'), api_key=api_key,safety_settings=safety_settings, temperature=0)


FEATURE_UPLOADS_BLOB_PREFIX = "feature_query_data_uploads/"

index = None
chat_engine = None


def load_documents_from_azure_fq(container_client: ContainerClient, prefix=FEATURE_UPLOADS_BLOB_PREFIX):
    documents = []
    blobs = container_client.list_blobs(name_starts_with=prefix)

    ALLOWED_EXTENSIONS = {".docx"}

    for blob in blobs:
        ext = os.path.splitext(blob.name)[1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            print(f"Skipping unsupported extension: {blob.name}")
            continue

        blob_client = container_client.get_blob_client(blob.name)
        try:
            raw_data = blob_client.download_blob().readall()
            if ext == ".docx":
                docx = DocxDocument(BytesIO(raw_data))
                text = "\n".join([para.text for para in docx.paragraphs])
            else:
                text = raw_data.decode('utf-8')

            documents.append(Document(text=text, metadata={"filename": blob.name}))

        except Exception as e:
            print(f"Skipping file due to error: {blob.name} — {e}")
            continue

    return documents



def build_index_fq(container_client, delete_old_index=True):
    global index, chat_engine

    print("Loading documents from Azure blob for reindexing Feature Query...")

    documents = load_documents_from_azure_fq(container_client)
    if not documents:
        print("No documents found in Azure blob folder for Feature Query, skipping index build.")
        index = None
        chat_engine = None
        return False

    if delete_old_index and os.path.exists(index_storage_dir):
        shutil.rmtree(index_storage_dir)

    index = VectorStoreIndex.from_documents(documents)
    index.storage_context.persist(persist_dir=index_storage_dir)

    chat_engine = index.as_query_engine()
    print("Index completed successfully for Feature Query.")
    return True 


def load_existing_index_fq(container_client):
    global index, chat_engine

    blobs = list(container_client.list_blobs(name_starts_with=FEATURE_UPLOADS_BLOB_PREFIX))
    if not blobs:
        print("No files found in Azure blob storage folder in Feature Query, skipping loading existing index.")
        index = None
        chat_engine = None
        return

    if os.path.exists(index_storage_dir) and os.listdir(index_storage_dir):
        try:
            print("Loading existing index from local storage for Feature Query...")
            storage_context = StorageContext.from_defaults(persist_dir=index_storage_dir)
            index = load_index_from_storage(storage_context)

            chat_engine = index.as_query_engine()
            print("Index and chat engine loaded successfully for Feature Query.")
        except Exception as e:
            print(f"Failed to load index for Feature Query: {e}")
            index = None
            chat_engine = None
    else:
        print("No existing local index found for Feature Query. Please run build_index() first.")
        index = None
        chat_engine = None


def get_feature_chat_engine():
    global chat_engine
    return chat_engine


def delete_blob_from_azure_fq(blob_name, container_client):
    try:
        full_blob_name = f"feature_query_data_uploads/{blob_name}"
        blob_client = container_client.get_blob_client(full_blob_name)
        blob_client.delete_blob()
        print(f"Deleted blob: {full_blob_name}")
        return True
    except Exception as e:
        print(f"Error deleting blob: {e}")
        return False

